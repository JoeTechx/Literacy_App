import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import re
import shutil
import os

if os.path.exists('Academic_Proposal.docx'):
    os.remove('Academic_Proposal.docx')

shutil.copy('OGUNDARE  thesis final 19 02 2026 correction again.docx', 'Academic_Proposal.docx')

doc = docx.Document('Academic_Proposal.docx')

for p in list(doc.paragraphs):
    p._element.getparent().remove(p._element)
for t in list(doc.tables):
    t._element.getparent().remove(t._element)

try:
    doc.styles['Normal'].paragraph_format.line_spacing = 2.0
except:
    pass


def add_centered_bold(text, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after:
        p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    return p


# ─────────────────────────────────────────────
# TITLE PAGE  (exactly mirrors OGUNDARE layout)
# ─────────────────────────────────────────────
TITLE = (
    "DESIGN AND DEVELOPMENT OF A MULTISENSORY MOBILE LEARNING APPLICATION "
    "FOR ENHANCING FOUNDATIONAL LITERACY SKILLS IN CHILDREN WITH DYSLEXIA "
    "IN ILORIN METROPOLIS"
)
add_centered_bold(TITLE, space_before=24)

add_centered_bold("M.Sc. Proposal", space_before=48)
add_centered_bold("By", space_before=48)
add_centered_bold("AGBI, John", space_before=24)
add_centered_bold("13/25PC156")

add_centered_bold("DEPARTMENT OF EDUCATIONAL TECHNOLOGY", space_before=48)
add_centered_bold("FACULTY OF EDUCATION")
add_centered_bold("UNIVERSITY OF ILORIN")
add_centered_bold("ILORIN, NIGERIA")

add_centered_bold("Supervisor: Dr. ____________________________", space_before=48)

# Venue / Date / Time (left-aligned, bold)
for label in [
    ("Venue: Departmental Postgraduate Seminar Room", 48),
    ("Date:\t__________________________________", 0),
    ("Time:\t10:00 am", 0),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if label[1]:
        p.paragraph_format.space_before = Pt(label[1])
    run = p.add_run(label[0])
    run.bold = True

# End of title page
doc.add_page_break()

# ─────────────────────────────────────────────
# REST OF DOCUMENT (from Academic_Proposal.md)
# ─────────────────────────────────────────────
with open('Academic_Proposal.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

# Find where TOC starts so we skip the old title page block
start_idx = 0
for i, line in enumerate(lines):
    if '**TABLE OF CONTENTS**' in line:
        start_idx = i
        break

for raw_line in lines[start_idx:]:
    line = raw_line.strip()
    if not line:
        continue

    if line == '---':
        continue

    # Chapter / References headings get a page break before them
    if line.startswith('# CHAPTER') or line.startswith('# REFERENCES'):
        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line[2:].upper())
        run.bold = True
        continue

    if line.startswith('## '):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(line[3:].upper())
        run.bold = True

    elif line.startswith('### '):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(line[4:])
        run.bold = True

    elif line.startswith('- ') or line.startswith('* '):
        try:
            p = doc.add_paragraph(style='List Bullet')
        except Exception:
            p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0
        text = line[2:]
        for part in re.split(r'(\*\*.*?\*\*)', text):
            if part.startswith('**') and part.endswith('**'):
                p.add_run(part[2:-2]).bold = True
            else:
                p.add_run(part)

    elif len(line) > 2 and line[0].isdigit() and line[1:3] == '. ':
        try:
            p = doc.add_paragraph(style='List Number')
        except Exception:
            p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0
        text = line[3:]
        for part in re.split(r'(\*\*.*?\*\*)', text):
            if part.startswith('**') and part.endswith('**'):
                p.add_run(part[2:-2]).bold = True
            else:
                p.add_run(part)

    else:
        if line == '**TABLE OF CONTENTS**':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('TABLE OF CONTENTS')
            run.bold = True
            doc.add_page_break()
            continue

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 2.0
        for part in re.split(r'(\*\*.*?\*\*)', line):
            if part.startswith('**') and part.endswith('**'):
                p.add_run(part[2:-2]).bold = True
            else:
                p.add_run(part)

doc.save('Academic_Proposal.docx')
print('Done – Academic_Proposal.docx saved successfully.')
